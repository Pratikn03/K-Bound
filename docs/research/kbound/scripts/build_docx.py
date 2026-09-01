#!/usr/bin/env python3
"""Build an editable, visually checked Word version of the compact K-Bound paper.

The maintained PDF is built from LaTeX with several constructs that Pandoc does
not resolve on its own: generated zero-argument macros, a custom figure macro,
manual numeric citations, and LaTeX cross-references.  This builder flattens the
live compact source, resolves those constructs without changing the manuscript,
converts the result to DOCX, and applies deterministic academic-paper styling.
"""

from __future__ import annotations

import argparse
import math
import re
import shutil
import subprocess
import tempfile
from collections.abc import Mapping, Sequence
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "kbound_submission.tex"
GENERATED_DIR = ROOT / "paper" / "generated"
FIGURE_ALTS = (
    "K-Bound decision flow from shadow candidate generation to adapt, freeze, or abstain.",
    "Population strict-commitment frontier over observable margin M and declared calibration-residual bound beta.",
    "KGA interval decisions for measured cell benefit; empirical residual coverage is not population-risk coverage.",
)

INPUT_PATTERN = re.compile(r"\\(?:input|include)\s*\{([^{}]+)\}")
GENERATED_MACRO_PATTERN = re.compile(
    r"\\(?:newcommand|renewcommand|providecommand)\s*"
    r"\{\s*\\(?P<name>[A-Za-z@]+)\s*\}\s*"
    r"(?:\[(?P<arity>\d+)\])?\s*\{"
)
SHA256_PATTERN = re.compile(r"[0-9a-fA-F]{64}")
HASH_MACRO_NAME_PATTERN = re.compile(r"(?:sha(?:256)?|hash)$", re.IGNORECASE)
CITATION_PATTERN = re.compile(
    r"\\cite(?:p|t|alp)?\*?(?:\s*\[[^]\r\n]*\]){0,2}\s*\{([^{}]+)\}",
    re.DOTALL,
)


def require_binary(name: str) -> str:
    path = shutil.which(name)
    if not path:
        raise RuntimeError(f"required executable not found: {name}")
    return path


def run_checked(command: list[str], *, cwd: Path, stdout_path: Path | None = None) -> str:
    if stdout_path is None:
        proc = subprocess.run(command, cwd=cwd, text=True, capture_output=True, check=False)
        if proc.returncode:
            raise RuntimeError(f"command failed ({proc.returncode}): {' '.join(command)}\n{proc.stderr}")
        return proc.stdout
    with stdout_path.open("w", encoding="utf-8") as handle:
        proc = subprocess.run(command, cwd=cwd, text=True, stdout=handle, stderr=subprocess.PIPE)
    if proc.returncode:
        raise RuntimeError(f"command failed ({proc.returncode}): {' '.join(command)}\n{proc.stderr}")
    return proc.stderr


def select_compact_conditionals(tex: str) -> str:
    """Select the named-author branch and exclude long-paper-only references."""
    named = re.sub(
        r"^\s*\\ifanon\b.*?^\s*\\else\b(.*?)^\s*\\fi\s*$",
        lambda match: match.group(1),
        tex,
        count=1,
        flags=re.DOTALL | re.MULTILINE,
    )
    named = re.sub(r"^\s*\\newif\\ifanon\s*$", "", named, flags=re.MULTILINE)
    named = re.sub(r"^\s*\\anonfalse\s*$", "", named, flags=re.MULTILINE)
    named = re.sub(
        r"\\ifdefined\\IncludeAuditableRefs\b.*?\\fi",
        "",
        named,
        flags=re.DOTALL,
    )
    return named


def strip_tex_comments(tex: str) -> str:
    """Remove comments while retaining escaped percent signs."""
    return "\n".join(re.sub(r"(?<!\\)%.*$", "", line) for line in tex.splitlines())


def _resolve_tex_input(owner: Path, target: str, search_root: Path) -> Path:
    """Resolve one literal ``\\input`` using the build root and owner directory."""
    if "\\" in target:
        raise RuntimeError(f"dynamic TeX input is not supported by the DOCX build: {target}")
    relative = Path(target.strip())
    if relative.suffix == "":
        relative = relative.with_suffix(".tex")
    candidates = [(search_root / relative).resolve(), (owner.parent / relative).resolve()]
    existing = list(dict.fromkeys(path for path in candidates if path.is_file()))
    if not existing:
        raise RuntimeError(f"referenced TeX input is missing: {target} (from {owner})")
    if len(existing) > 1:
        raise RuntimeError(f"ambiguous TeX input {target} from {owner}: " + ", ".join(str(path) for path in existing))
    return existing[0]


def discover_generated_macro_sources(
    source: Path = SOURCE,
    *,
    generated_dir: Path = GENERATED_DIR,
    search_root: Path | None = None,
) -> tuple[Path, ...]:
    """Find every referenced ``*_numbers.tex`` file in the live source graph.

    Merely placing a prospective result file in ``paper/generated`` does not
    activate it.  The maintained manuscript must reference it through its TeX
    dependency graph.  Conversely, a referenced missing or malformed file is a
    hard failure rather than a silent fallback to stale values.
    """
    source = source.resolve()
    generated_dir = generated_dir.resolve()
    search_root = (search_root or source.parent).resolve()
    if not source.is_file():
        raise RuntimeError(f"DOCX source is missing: {source}")

    visiting: set[Path] = set()
    visited: set[Path] = set()
    generated_sources: list[Path] = []

    def visit(path: Path) -> None:
        if path in visiting:
            raise RuntimeError(f"cyclic TeX input detected at {path}")
        if path in visited:
            return
        visiting.add(path)
        text = strip_tex_comments(path.read_text(encoding="utf-8"))
        for raw_target in INPUT_PATTERN.findall(text):
            dependency = _resolve_tex_input(path, raw_target, search_root)
            if dependency.is_relative_to(generated_dir) and dependency.name.endswith("_numbers.tex"):
                if dependency not in generated_sources:
                    generated_sources.append(dependency)
            visit(dependency)
        visiting.remove(path)
        visited.add(path)

    visit(source)
    if not generated_sources:
        raise RuntimeError("live manuscript references no generated *_numbers.tex source")
    return tuple(generated_sources)


def _zero_argument_macros(path: Path) -> dict[str, str]:
    """Parse braced zero-argument command definitions, including nested values."""
    text = strip_tex_comments(path.read_text(encoding="utf-8"))
    macros: dict[str, str] = {}
    position = 0
    while match := GENERATED_MACRO_PATTERN.search(text, position):
        opening = match.end() - 1
        closing = find_brace_end(text, opening)
        position = closing + 1
        if match.group("arity") not in (None, "0"):
            continue
        name = match.group("name")
        value = text[opening + 1 : closing].strip()
        if not value:
            raise RuntimeError(f"generated macro {name} has an empty value in {path}")
        if name in macros:
            raise RuntimeError(f"generated macro {name} is defined more than once in {path}")
        macros[name] = value
    if not macros:
        raise RuntimeError(f"generated macro source defines no zero-argument commands: {path}")
    return macros


def generated_macros(sources: Sequence[Path] | None = None) -> dict[str, str]:
    """Load a conflict-free macro map from all live generated number sources."""
    source_paths = tuple(discover_generated_macro_sources() if sources is None else sources)
    if not source_paths:
        raise RuntimeError("no generated macro sources were supplied")
    macros: dict[str, str] = {}
    origins: dict[str, Path] = {}
    for path in source_paths:
        for name, value in _zero_argument_macros(path).items():
            if name in macros:
                raise RuntimeError(f"generated macro {name} is defined by both {origins[name]} and {path}")
            macros[name] = value
            origins[name] = path
    return _resolve_generated_macro_values(macros)


def _resolve_generated_macro_values(macros: Mapping[str, str]) -> dict[str, str]:
    """Resolve references among generated zero-argument macros deterministically.

    Generated sources are allowed to factor a value through another generated
    command.  A one-pass body substitution is order-dependent and can leave a
    transitive command unresolved, including a provenance hash.  Resolve the
    generated dependency graph first and fail closed on cycles; ordinary LaTeX
    commands that are not generated values remain untouched.
    """

    values = dict(macros)
    resolved: dict[str, str] = {}
    visiting: list[str] = []

    def resolve(name: str) -> str:
        if name in resolved:
            return resolved[name]
        if name in visiting:
            cycle = " -> ".join((*visiting[visiting.index(name) :], name))
            raise RuntimeError(f"cyclic generated macro definition: {cycle}")
        visiting.append(name)
        value = values[name]

        def replace(match: re.Match[str]) -> str:
            dependency = match.group(1)
            return resolve(dependency) if dependency in values else match.group(0)

        value = re.sub(r"\\([A-Za-z@]+)(?![A-Za-z@])", replace, value)
        visiting.pop()
        resolved[name] = value
        return value

    for macro_name in values:
        resolve(macro_name)
    return resolved


def _strip_body_generated_macro_definitions(
    body: str,
    macro_names: Sequence[str],
) -> str:
    """Remove harvested generated definitions before body-value expansion.

    ``latexpand`` may inline a live ``*_numbers.tex`` input after
    ``\\begin{document}``.  Those definitions are redundant once their values
    have been harvested, and substituting the command token inside its own
    declaration would turn ``\\newcommand{\\Count}{5}`` into the invalid
    ``\\newcommand{5}{5}``.  Remove only recognized zero-argument generated
    definitions; preserve every ordinary manuscript command definition.
    """

    recognized = set(macro_names)
    pieces: list[str] = []
    cursor = 0
    search_from = 0
    while match := GENERATED_MACRO_PATTERN.search(body, search_from):
        opening = match.end() - 1
        closing = find_brace_end(body, opening)
        search_from = closing + 1
        if match.group("arity") not in (None, "0") or match.group("name") not in recognized:
            continue
        pieces.append(body[cursor : match.start()])
        cursor = closing + 1
    pieces.append(body[cursor:])
    return "".join(pieces)


def generated_value_requirements(tex: str, macros: Mapping[str, str]) -> dict[str, int]:
    """Derive document-value checks from generated macros used in the body.

    Hash-like macros are deliberately strict: once the manuscript uses one, its
    generated value must be a complete SHA-256 and must survive conversion at
    least as many times as the source invokes it.
    """
    marker = r"\begin{document}"
    if marker not in tex:
        raise RuntimeError("flattened source has no document body")
    _, body = tex.split(marker, 1)
    body = strip_tex_comments(body)
    values = _resolve_generated_macro_values(macros)
    body = _strip_body_generated_macro_definitions(body, tuple(values))
    requirements: dict[str, int] = {}
    for name, value in values.items():
        count = len(re.findall(rf"\\{re.escape(name)}(?![A-Za-z@])", body))
        if not count:
            continue
        if HASH_MACRO_NAME_PATTERN.search(name) and not SHA256_PATTERN.fullmatch(value):
            raise RuntimeError(f"used generated hash macro {name} is not a complete SHA-256 value")
        if SHA256_PATTERN.fullmatch(value):
            normalized = value.lower()
            requirements[normalized] = requirements.get(normalized, 0) + count
    return requirements


def expand_generated_macros(tex: str, macros: Mapping[str, str] | None = None) -> str:
    marker = r"\begin{document}"
    if marker not in tex:
        raise RuntimeError("flattened source has no document body")
    preamble, body = tex.split(marker, 1)
    values = _resolve_generated_macro_values(generated_macros() if macros is None else macros)
    body = _strip_body_generated_macro_definitions(body, tuple(values))
    for name, value in sorted(values.items(), key=lambda item: -len(item[0])):
        body = re.sub(
            rf"\\{re.escape(name)}(?![A-Za-z@])",
            lambda _, replacement=value: replacement,
            body,
        )
    return preamble + marker + body


def unwrap_breakable_hashes(tex: str) -> str:
    """Keep the hash text while removing the PDF-only line-breaking wrapper."""
    return re.sub(r"\\hashtext\{([^{}]+)\}", r"\1", tex)


def replace_actions_in_math(tex: str) -> str:
    replacements = {
        r"\adapt": r"\text{adapt}",
        r"\freeze": r"\text{freeze}",
        r"\abstain": r"\text{abstain}",
        r"\textsc{adapt}": r"\text{adapt}",
        r"\textsc{freeze}": r"\text{freeze}",
        r"\textsc{abstain}": r"\text{abstain}",
    }

    def repair(fragment: str) -> str:
        for source, target in replacements.items():
            fragment = fragment.replace(source, target)
        return fragment

    env_pattern = re.compile(r"\\begin\{(equation\*?|align\*?)\}(.*?)\\end\{\1\}", re.DOTALL)
    tex = env_pattern.sub(
        lambda match: rf"\begin{{{match.group(1)}}}" + repair(match.group(2)) + rf"\end{{{match.group(1)}}}",
        tex,
    )
    tex = re.sub(
        r"(?<!\\)\\\[(.*?)(?<!\\)\\\]",
        lambda match: r"\[" + repair(match.group(1)) + r"\]",
        tex,
        flags=re.DOTALL,
    )
    tex = re.sub(
        r"(?<!\\)\$(?!\$)(.*?)(?<!\\)\$",
        lambda match: "$" + repair(match.group(1)) + "$",
        tex,
        flags=re.DOTALL,
    )
    return tex


def expand_text_action_words(tex: str) -> str:
    """Expand action macros that Pandoc otherwise drops in ordinary prose."""
    marker = r"\begin{document}"
    preamble, body = tex.split(marker, 1)
    for macro, word in (("adapt", "ADAPT"), ("freeze", "FREEZE"), ("abstain", "ABSTAIN")):
        body = re.sub(rf"\\{macro}(?![A-Za-z@])", word, body)
        body = body.replace(rf"\textsc{{{macro}}}", word)
    return preamble + marker + body


def simplify_alignment_math(tex: str) -> str:
    """Remove alignment markers that Word exposes as literal ampersands."""

    def split_align(match: re.Match[str]) -> str:
        lines = [line.strip().replace("&", "") for line in re.split(r"\\\\", match.group(2))]
        rendered = []
        for line in lines:
            if not line:
                continue
            suppressed = re.search(r"\\(?:notag|nonumber)\b", line) is not None
            line = re.sub(r"\\(?:notag|nonumber)\b", "", line)
            if match.group(1) == "align" and not suppressed:
                # Preserve each numbered row and label for the shared equation
                # counter. Turning align into \[...\] loses cross-references.
                rendered.append(r"\begin{equation}" + line + r"\end{equation}")
            else:
                rendered.append(r"\[" + line + r"\]")
        return "\n".join(rendered)

    tex = re.sub(
        r"\\begin\{(align\*?)\}(.*?)\\end\{\1\}",
        split_align,
        tex,
        flags=re.DOTALL,
    )

    def flatten_aligned(match: re.Match[str]) -> str:
        lines = [line.strip().replace("&", "") for line in re.split(r"\\\\", match.group(1))]
        return r"\quad".join(line for line in lines if line)

    tex = re.sub(
        r"\\begin\{aligned\}(.*?)\\end\{aligned\}",
        flatten_aligned,
        tex,
        flags=re.DOTALL,
    )
    return re.sub(
        r"\\underbrace\{([^{}]+)\}_\{\\text\{[^{}]+\}\}",
        lambda match: match.group(1),
        tex,
    )


def replace_algorithm_for_word(tex: str) -> str:
    """Translate the manuscript's algorithmic subset into editable Word lists.

    Statements and conditions come from the source, not a second algorithm.
    Unknown commands and malformed branches must fail before Pandoc can discard
    an instruction. Nested enumerations preserve each conditional branch.
    """
    pattern = re.compile(
        r"\\begin\{algorithm\}(?:\[[^]\r\n]*\])?(.*?)\\end\{algorithm\}",
        re.DOTALL,
    )
    matches = list(pattern.finditer(tex))
    if len(matches) != 1:
        raise RuntimeError(f"expected one KGA algorithm, found {len(matches)}")
    match = matches[0]
    content = strip_tex_comments(match.group(1)).strip()
    bodies = list(re.finditer(
        r"\\begin\{algorithmic\}(?:\[\d+\])?(.*?)\\end\{algorithmic\}",
        content,
        flags=re.DOTALL,
    ))
    if len(bodies) != 1:
        raise RuntimeError("expected one algorithmic body in the KGA algorithm")
    body = bodies[0]
    controls = {"Require", "State", "If", "ElsIf", "Else", "EndIf"}
    inline_commands = {
        "mathcal", "phi", "theta", "leftarrow", "widehat", "Delta", "varepsilon",
        "eqref", "ref", "text", "textbf", "texttt", "textsc", "emph", "mathrm",
        "adapt", "freeze", "abstain",
    }

    def validate_inline(value: str) -> str:
        unknown = set(re.findall(r"\\([A-Za-z@]+)", value)) - inline_commands
        if unknown:
            raise RuntimeError("unsupported algorithm text command(s): " + ", ".join(sorted(unknown)))
        return " ".join(value.split())

    def argument(value: str, context: str) -> tuple[str, str]:
        value = value.lstrip()
        if not value.startswith("{"):
            raise RuntimeError(f"algorithm {context} requires a literal braced argument")
        try:
            closing = find_brace_end(value, 0)
        except RuntimeError as error:
            raise RuntimeError(f"unbalanced algorithm {context} argument") from error
        result = value[1:closing].strip()
        if not result:
            raise RuntimeError(f"empty algorithm {context} argument")
        return result, value[closing + 1:]

    metadata = (content[:body.start()] + content[body.end():]).strip()
    caption = None
    labels = []
    while metadata:
        command = re.match(r"\\(caption|label)\b", metadata)
        if command is None:
            raise RuntimeError("unsupported algorithm metadata: " + metadata[:80])
        value, metadata = argument(metadata[command.end():], command.group(1))
        metadata = metadata.strip()
        if command.group(1) == "caption":
            if caption is not None:
                raise RuntimeError("expected one algorithm caption")
            caption = validate_inline(value)
        else:
            if not re.fullmatch(r"[A-Za-z0-9:._-]+", value):
                raise RuntimeError("unsupported algorithm label: " + value)
            labels.append(r"\label{" + value + "}")
    if caption is None:
        raise RuntimeError("expected one algorithm caption")

    source = body.group(1).strip()
    unknown = set(re.findall(r"\\([A-Za-z@]+)", source)) - controls - inline_commands
    if unknown:
        raise RuntimeError("unsupported algorithm command(s): " + ", ".join(sorted(unknown)))
    tokens = list(re.finditer(r"\\(Require|State|If|ElsIf|Else|EndIf)\b", source))
    if not tokens or source[:tokens[0].start()].strip():
        raise RuntimeError("unsupported algorithm text before its first statement")
    lines = [r"\paragraph{Algorithm 1. " + caption + "}", *labels, r"\begin{enumerate}"]
    branches = []
    state_count = 0
    for index, token in enumerate(tokens):
        end = tokens[index + 1].start() if index + 1 < len(tokens) else len(source)
        value = source[token.end():end].strip()
        command = token.group(1)
        if command in {"Require", "State"}:
            if not value:
                raise RuntimeError(f"empty algorithm {command} statement")
            value = validate_inline(value)
            if command == "Require":
                if index != 0:
                    raise RuntimeError("algorithm Require must precede all action statements")
                lines.append(r"\item \textbf{Inputs:} " + value)
            else:
                state_count += 1
                lines.append(r"\item " + value)
                if branches:
                    branches[-1]["has_content"] = True
            continue
        if command in {"If", "ElsIf"}:
            condition, remainder = argument(value, command)
            if remainder.strip():
                raise RuntimeError("unsupported algorithm text after a condition")
            condition = validate_inline(condition)
        elif value:
            raise RuntimeError(f"unsupported algorithm text after {command}")
        if command == "If":
            if branches:
                branches[-1]["has_content"] = True
            lines.extend((r"\item \textbf{If} " + condition + ":", r"\begin{enumerate}"))
            branches.append({"has_else": False, "has_content": False})
        else:
            if not branches or not branches[-1]["has_content"]:
                raise RuntimeError(f"algorithm {command} has no preceding nonempty branch")
            lines.append(r"\end{enumerate}")
            if command == "EndIf":
                branches.pop()
                continue
            if branches[-1]["has_else"]:
                raise RuntimeError(f"algorithm {command} cannot follow Else")
            branches[-1]["has_else"] = command == "Else"
            branches[-1]["has_content"] = False
            heading = r"\textbf{Else}:" if command == "Else" else r"\textbf{Else if} " + condition + ":"
            lines.extend((r"\item " + heading, r"\begin{enumerate}"))
    if branches:
        raise RuntimeError("algorithm has an unclosed If branch")
    if tokens[0].group(1) != "Require" or not state_count:
        raise RuntimeError("algorithm requires inputs and at least one action statement")
    lines.append(r"\end{enumerate}")
    return tex[:match.start()] + "\n".join(lines) + tex[match.end():]


def normalize_word_theory_notation(tex: str) -> str:
    """Preserve the overbar and explicit Roman theorem clauses in Word.

    Pandoc emits an accent for ``\\bar a`` that LibreOffice can display as an
    acute mark. The equivalent ``\\overline{a}`` emits an editable OMML top bar.
    Pandoc also drops the source's manual ``\\item[\\textnormal{(i)}]`` labels;
    its native Roman-list syntax keeps those labels aligned with the proof.
    These conversions affect the Word input only, not the maintained TeX/PDF.
    """
    tex = re.sub(r"\\bar(?:\s*\{\s*a\s*\}|\s+a)(?![A-Za-z@])", r"\\overline{a}", tex)
    roman_item = re.compile(r"\\item\[\\textnormal\{\(([ivxlcdm]+)\)\}\]")

    def normalize_list(match: re.Match[str]) -> str:
        content = match.group(1)
        labels = roman_item.findall(content)
        if not labels:
            return match.group(0)
        # The maintained theorem has four clauses. Fail visibly if that source
        # contract changes instead of silently dropping or relabeling items.
        expected = ["i", "ii", "iii", "iv"]
        if labels != expected or len(re.findall(r"\\item\b", content)) != len(expected):
            raise RuntimeError("expected four consecutive Roman theorem clauses (i)--(iv)")
        return r"\begin{enumerate}[(i)]" + roman_item.sub(r"\\item", content) + r"\end{enumerate}"

    return re.sub(
        r"\\begin\{enumerate\}(.*?)\\end\{enumerate\}", normalize_list, tex, flags=re.DOTALL
    )


def replace_figure_macro(tex: str) -> str:
    pattern = re.compile(r"\\kbgraphics(?:\[[^\]]*\])?\{([^}]+)\}")
    return pattern.sub(r"\\includegraphics[width=0.80\\textwidth]{\1}", tex)


def unwrap_resizebox_tables(tex: str) -> str:
    """Expose tabular content hidden inside PDF-only resizebox wrappers."""
    pattern = re.compile(
        r"\\resizebox\{[^}]+\}\{!\}\{(\s*\\begin\{tabular\}.*?\\end\{tabular\})\s*\}",
        flags=re.DOTALL,
    )
    return pattern.sub(lambda match: match.group(1), tex)


def normalize_starred_floats(tex: str) -> str:
    """Word is single-column, so preserve starred-float captions as normal floats."""
    return tex.replace(r"\begin{table*}", r"\begin{table}").replace(r"\end{table*}", r"\end{table}")


def split_bibliography(tex: str) -> tuple[str, list[tuple[str, str]]]:
    matches = list(
        re.finditer(
            r"\\begin\{thebibliography\}\{[^}]*\}(.*?)\\end\{thebibliography\}",
            tex,
            flags=re.DOTALL,
        )
    )
    if not matches:
        raise RuntimeError("active compact bibliography not found")
    if len(matches) != 1:
        raise RuntimeError(f"expected one active compact bibliography, found {len(matches)}")
    match = matches[0]
    raw = match.group(1)
    # The shared bibliography uses a three-argument wrapper so the TMLR driver
    # can emit author-year labels while the compact PDF remains numeric.  Word
    # output is numeric, so collapse each wrapper invocation to an ordinary
    # \bibitem before extracting entries.  The metadata arguments are flat by
    # construction; reject nested braces by leaving an unrecognized invocation
    # for the fail-closed empty-bibliography check below.
    raw = re.sub(
        r"\\KBbibitem\{([^{}]+)\}\{[^{}]*\}\{[^{}]*\}",
        r"\\bibitem{\1}",
        raw,
    )
    item_pattern = re.compile(r"\\bibitem(?:\[[^\]]*\])?\{([^}]+)\}")
    markers = list(item_pattern.finditer(raw))
    entries: list[tuple[str, str]] = []
    for index, marker in enumerate(markers):
        end = markers[index + 1].start() if index + 1 < len(markers) else len(raw)
        content = raw[marker.end() : end].strip()
        content = re.sub(r"\\newblock\s*", " ", content)
        if not content:
            raise RuntimeError(f"active compact bibliography entry is empty: {marker.group(1)}")
        entries.append((marker.group(1), content))
    if not entries:
        raise RuntimeError(
            "active compact bibliography has no references; expected \\bibitem or "
            "flat \\KBbibitem entries"
        )
    keys = [key for key, _ in entries]
    duplicate_keys = sorted({key for key in keys if keys.count(key) > 1})
    if duplicate_keys:
        raise RuntimeError("active compact bibliography has duplicate key(s): " + ", ".join(duplicate_keys))
    rendered = [r"\section*{References}", r"\begin{enumerate}"]
    rendered.extend(r"\item " + content for _, content in entries)
    rendered.append(r"\end{enumerate}")
    return tex[: match.start()] + "\n".join(rendered) + tex[match.end() :], entries


def compress_numbers(numbers: list[int]) -> str:
    ordered = sorted(dict.fromkeys(numbers))
    groups: list[str] = []
    start = previous = ordered[0]
    for value in ordered[1:] + [math.inf]:
        if value == previous + 1:
            previous = int(value)
            continue
        if start == previous:
            groups.append(str(start))
        elif previous == start + 1:
            groups.extend((str(start), str(previous)))
        else:
            groups.append(f"{start}--{previous}")
        if value is not math.inf:
            start = previous = int(value)
    return ", ".join(groups)


def resolve_citations(tex: str, entries: list[tuple[str, str]]) -> str:
    numbers = {key: index for index, (key, _) in enumerate(entries, start=1)}

    def replace(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).replace("\n", "").split(",")]
        missing = [key for key in keys if key not in numbers]
        if missing:
            raise RuntimeError(f"unknown citation key(s): {', '.join(missing)}")
        return "[" + compress_numbers([numbers[key] for key in keys]) + "]"

    tex = CITATION_PATTERN.sub(replace, tex)
    unresolved = sorted(set(re.findall(r"\\cite[A-Za-z@*]*", tex)))
    if unresolved:
        raise RuntimeError("unsupported or unresolved citation command(s): " + ", ".join(unresolved))
    return tex


def find_brace_end(text: str, opening: int) -> int:
    depth = 0
    for index in range(opening, len(text)):
        if text[index] == "{" and (index == 0 or text[index - 1] != "\\"):
            depth += 1
        elif text[index] == "}" and (index == 0 or text[index - 1] != "\\"):
            depth -= 1
            if depth == 0:
                return index
    raise RuntimeError("unbalanced caption braces")


def prefix_caption(content: str, prefix: str) -> str:
    start = content.find(r"\caption{")
    if start < 0:
        return content
    opening = start + len(r"\caption")
    closing = find_brace_end(content, opening)
    caption = content[opening + 1 : closing]
    if caption.lstrip().startswith(prefix):
        return content
    return content[: opening + 1] + prefix + caption + content[closing:]


def number_environments(tex: str) -> tuple[str, dict[str, str]]:
    labels: dict[str, str] = {}

    def process_family(source: str, expression: str, family: str) -> str:
        pattern = re.compile(rf"\\begin\{{({expression})\}}(.*?)\\end\{{\1\}}", re.DOTALL)
        counter = 0

        def replace(match: re.Match[str]) -> str:
            nonlocal counter
            counter += 1
            environment = match.group(1)
            content = match.group(2)
            for label in re.findall(r"\\label\{([^}]+)\}", content):
                labels[label] = str(counter)
            content = re.sub(r"\\label\{[^}]+\}", "", content)
            if family == "figure":
                content = prefix_caption(content, f"Figure {counter}. ")
            elif family == "table":
                content = prefix_caption(content, f"Table {counter}. ")
            elif family == "algorithm":
                content = prefix_caption(content, f"Algorithm {counter}. ")
            elif family == "equation":
                content = re.sub(r"[.,]\s*$", "", content.rstrip())
                content = content + rf"\quad\text{{({counter})}}" + "\n"
            return rf"\begin{{{environment}}}" + content + rf"\end{{{environment}}}"

        return pattern.sub(replace, source)

    # Process inner display environments before theorem-like containers so a
    # theorem that contains an equation cannot hide or renumber that equation.
    tex = process_family(tex, r"figure\*?", "figure")
    tex = process_family(tex, r"table\*?", "table")
    tex = process_family(tex, r"equation", "equation")
    tex = process_family(tex, r"algorithm", "algorithm")
    for theorem_family in (
        "theorem",
        "lemma",
        "proposition",
        "corollary",
        "definition",
        "assumption",
        "remark",
    ):
        tex = process_family(tex, theorem_family, theorem_family)
    return tex, labels


def number_section_headings(tex: str) -> tuple[str, dict[str, str]]:
    """Materialize article-style heading numbers and their cross-reference map.

    Pandoc does not number these Word headings by default.  Use the same
    source-derived numbers in the visible headings and references, just as for
    captions and equations above.  Every unstarred heading advances its counter,
    whether or not it has a label; ``\\appendix`` resets sections to A, B, ... .
    Labels may be inside the title or immediately after it (including comments).
    This is a literal-heading exporter, not a general TeX counter interpreter.
    """
    # Preserve source offsets while preventing commented-out commands or braces
    # from participating in the heading scan.
    scan = re.sub(r"(?<!\\)%[^\n]*", lambda match: " " * len(match.group()), tex)
    pattern = re.compile(r"(?<!\\)\\(appendix|section|subsection|subsubsection)(?![A-Za-z@])")
    label_pattern = re.compile(r"\s*\\label\s*\{([^{}]+)\}")
    labels: dict[str, str] = {}
    counters = [0, 0, 0]
    appendix = False
    insertions: list[tuple[int, str]] = []
    position = 0

    while match := pattern.search(scan, position):
        command = match.group(1)
        position = match.end()
        if command == "appendix":
            counters = [0, 0, 0]
            appendix = True
            continue

        while position < len(scan) and scan[position].isspace():
            position += 1
        starred = position < len(scan) and scan[position] == "*"
        if starred:
            position += 1
        while position < len(scan) and scan[position].isspace():
            position += 1
        if position < len(scan) and scan[position] == "[":
            # An optional short title ends at an unbraced closing bracket.
            depth = 0
            position += 1
            while position < len(scan):
                char = scan[position]
                if position == 0 or scan[position - 1] != "\\":
                    if char == "{":
                        depth += 1
                    elif char == "}":
                        depth -= 1
                    elif char == "]" and depth == 0:
                        break
                position += 1
            if position == len(scan):
                raise RuntimeError(f"unbalanced optional title for {command}")
            position += 1
            while position < len(scan) and scan[position].isspace():
                position += 1
        if position == len(scan) or scan[position] != "{":
            raise RuntimeError(f"expected a literal braced title for {command}")
        opening = position
        try:
            closing = find_brace_end(scan, opening)
        except RuntimeError as error:
            raise RuntimeError(f"unbalanced title for {command}") from error
        position = closing + 1
        if starred:
            continue

        level = ("section", "subsection", "subsubsection").index(command)
        counters[level] += 1
        counters[level + 1 :] = [0] * (2 - level)
        if appendix:
            if not 1 <= counters[0] <= 26:
                raise RuntimeError("appendix headings require an article-style section from A to Z")
            section = chr(ord("A") + counters[0] - 1)
        else:
            section = str(counters[0])
        number = ".".join([section, *(str(value) for value in counters[1 : level + 1])])
        insertions.append((opening + 1, number + " "))

        heading_labels = re.findall(r"\\label\s*\{([^{}]+)\}", scan[opening + 1 : closing])
        while label_match := label_pattern.match(scan, position):
            heading_labels.append(label_match.group(1))
            position = label_match.end()
        for label in heading_labels:
            if label in labels:
                raise RuntimeError(f"duplicate section label: {label}")
            labels[label] = number

    for offset, prefix in reversed(insertions):
        tex = tex[:offset] + prefix + tex[offset:]
    return tex, labels


def section_labels(tex: str) -> dict[str, str]:
    """Return the same label map used to number the visible Word headings."""
    return number_section_headings(tex)[1]


def resolve_cross_references(tex: str) -> str:
    tex, labels = number_environments(tex)
    tex, heading_labels = number_section_headings(tex)
    labels.update(heading_labels)

    def eqref(match: re.Match[str]) -> str:
        label = match.group(1)
        if label not in labels:
            raise RuntimeError(f"unresolved equation reference: {label}")
        return f"({labels[label]})"

    def ref(match: re.Match[str]) -> str:
        label = match.group(1)
        if label not in labels:
            raise RuntimeError(f"unresolved reference: {label}")
        return labels[label]

    tex = re.sub(r"\\eqref\{([^}]+)\}", eqref, tex)
    tex = re.sub(r"\\ref\{([^}]+)\}", ref, tex)
    tex = re.sub(r"\\label\{[^}]+\}", "", tex)
    if re.search(r"\\(?:eqref|ref)\{", tex):
        raise RuntimeError("unresolved cross-reference command remains")
    return tex


def preprocess_with_metadata(tex: str, *, macros: Mapping[str, str] | None = None) -> tuple[str, int, dict[str, int]]:
    """Flatten conversion-only constructs and return data-derived build checks."""
    values = _resolve_generated_macro_values(generated_macros() if macros is None else macros)
    tex = select_compact_conditionals(tex)
    required_values = generated_value_requirements(tex, values)
    tex = expand_generated_macros(tex, values)
    tex = unwrap_breakable_hashes(tex)
    tex = replace_actions_in_math(tex)
    tex = expand_text_action_words(tex)
    tex = simplify_alignment_math(tex)
    tex = normalize_word_theory_notation(tex)
    tex = replace_algorithm_for_word(tex)
    tex = replace_figure_macro(tex)
    tex = unwrap_resizebox_tables(tex)
    tex = normalize_starred_floats(tex)
    tex, entries = split_bibliography(tex)
    tex = resolve_citations(tex, entries)
    tex = resolve_cross_references(tex)
    return tex, len(entries), required_values


def preprocess(tex: str) -> str:
    """Compatibility wrapper returning only the preprocessed TeX text."""
    return preprocess_with_metadata(tex)[0]


def set_font(style, name: str, size: float, *, bold: bool | None = None, italic: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attribute}"), name)


def set_style_spacing(style, *, before: float, after: float, line: float) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_styles(doc: Document) -> None:
    body_names = ("Normal", "Body Text", "First Paragraph")
    for name in body_names:
        if name in doc.styles:
            set_font(doc.styles[name], "Times New Roman", 9.8)
            set_style_spacing(doc.styles[name], before=0, after=4, line=1.08)
    specifications = {
        "Title": (17.0, True, False, 0, 5),
        "Author": (9.5, False, False, 0, 7),
        "Abstract Title": (10.0, True, False, 5, 2),
        "Abstract": (9.3, False, False, 0, 6),
        "Heading 1": (13.0, True, False, 11, 4),
        "Heading 2": (11.0, True, False, 9, 3),
        "Heading 3": (10.2, True, False, 7, 2),
        "Heading 4": (9.8, True, False, 5, 1),
        "Image Caption": (8.4, False, True, 2, 6),
        "Table Caption": (8.4, False, True, 5, 3),
        "Caption": (8.4, False, True, 3, 5),
        "Source Code": (8.2, False, False, 2, 4),
    }
    for name, (size, bold, italic, before, after) in specifications.items():
        if name not in doc.styles:
            continue
        font = "Courier New" if name == "Source Code" else "Times New Roman"
        set_font(doc.styles[name], font, size, bold=bold, italic=italic)
        set_style_spacing(doc.styles[name], before=before, after=after, line=1.0 if size < 9 else 1.05)
    for name in ("Title", "Author", "Abstract Title", "Image Caption", "Table Caption", "Caption"):
        if name in doc.styles:
            doc.styles[name].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "Abstract" in doc.styles:
        doc.styles["Abstract"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        if name in doc.styles:
            doc.styles[name].paragraph_format.keep_with_next = True
    # Table captions precede their tables; image captions can follow an image.
    if "Table Caption" in doc.styles:
        doc.styles["Table Caption"].paragraph_format.keep_with_next = True


def normalize_compact_heading_hierarchy(doc: Document) -> None:
    """Map LaTeX paragraph headings to valid Word heading levels.

    Pandoc maps every LaTeX ``\\paragraph`` to Heading 4, even when the
    manuscript has no intervening Heading 2 or Heading 3.  Keep the compact
    visual treatment while assigning Heading 2 below a section and Heading 3
    below a subsection so assistive technology receives a coherent outline.
    """
    parent_level = 1
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        match = re.fullmatch(r"Heading ([123])", style_name)
        if match:
            parent_level = int(match.group(1))
            continue
        if style_name != "Heading 4":
            continue
        target_level = 2 if parent_level == 1 else 3
        paragraph.style = doc.styles[f"Heading {target_level}"]
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.05
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9.8)
            run.font.bold = True
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                rfonts.set(qn(f"w:{attribute}"), "Times New Roman")


def set_cell_margins(cell, *, top: int = 60, start: int = 80, bottom: int = 60, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7B7B7")


def table_widths(table, total: int) -> list[int]:
    headers = tuple(
        re.sub(r"\s+", " ", cell.text).strip().lower()
        for cell in table.rows[0].cells
    )
    if headers == ("track", "shift geometry", "candidate/protocol", "defensible use"):
        return [2150, 2300, 2200, total - 6650]
    if headers == ("track", "released status", "present evidence", "remaining limitation"):
        return [1800, 1600, 3100, total - 6500]

    columns = len(table.columns)
    weights: list[float] = []
    for column in range(columns):
        lengths = [len(row.cells[column].text.strip()) for row in table.rows]
        nonempty = [value for value in lengths if value]
        peak = max(nonempty, default=8)
        average = sum(nonempty) / len(nonempty) if nonempty else 8
        weights.append(max(7.0, min(44.0, 0.38 * peak + 0.62 * average)))
    widths = [max(560, round(total * weight / sum(weights))) for weight in weights]
    difference = total - sum(widths)
    widths[-1] += difference
    if widths[-1] < 560:
        deficit = 560 - widths[-1]
        widths[-1] = 560
        widest = max(range(len(widths) - 1), key=widths.__getitem__)
        widths[widest] -= deficit
    return widths


def protect_numeric_citations_in_tables(doc: Document) -> None:
    """Keep bracketed numeric citations intact in narrow Word table cells."""

    pattern = re.compile(r"\[([0-9]+(?:--[0-9]+)?(?:,\s*[0-9]+(?:--[0-9]+)?)*)\]")
    word_joiner = "\u2060"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = pattern.sub(
                            lambda match: (
                                f"[{word_joiner}"
                                + match.group(1).replace(" ", "\u00a0")
                                + f"{word_joiner}]"
                            ),
                            run.text,
                        )


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table._tbl.tr_lst:
        column = 0
        for cell in row.tc_lst:
            tc_pr = cell.get_or_add_tcPr()
            span_node = tc_pr.first_child_found_in("w:gridSpan")
            span = int(span_node.get(qn("w:val"))) if span_node is not None else 1
            width = sum(widths[column : column + span])
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            column += span


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        columns = len(table.columns)
        keep_table_together = len(table.rows) <= 7
        widths = table_widths(table, 9960)
        set_table_geometry(table, widths)
        set_table_borders(table)
        size = 7.4 if columns >= 7 else 7.8 if columns >= 5 else 8.3 if columns == 4 else 8.8
        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = tr_pr.find(qn("w:cantSplit"))
            if cant_split is None:
                cant_split = OxmlElement("w:cantSplit")
                tr_pr.append(cant_split)
            cant_split.set(qn("w:val"), "true")
            if row_index == 0:
                repeat = tr_pr.find(qn("w:tblHeader"))
                if repeat is None:
                    repeat = OxmlElement("w:tblHeader")
                    tr_pr.append(repeat)
                repeat.set(qn("w:val"), "true")
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shade = tc_pr.first_child_found_in("w:shd")
                    if shade is None:
                        shade = OxmlElement("w:shd")
                        tc_pr.append(shade)
                    shade.set(qn("w:fill"), "E8EEF5")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    # Keep a header with its first complete body row.  Longer
                    # tables may paginate between body rows, never within one.
                    paragraph.paragraph_format.keep_with_next = (
                        row_index < len(table.rows) - 1 and (row_index == 0 or keep_table_together)
                    )
                    text = paragraph.text.strip()
                    if row_index == 0 or re.fullmatch(r"[-+()0-9.,/%\s]+", text or "x"):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(size)
                        if row_index == 0:
                            run.font.bold = True
                        if run._element.rPr is not None:
                            rfonts = run._element.rPr.rFonts
                            if rfonts is not None:
                                rfonts.set(qn("w:ascii"), "Times New Roman")
                                rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, value, end):
        run._r.append(node)


def bibliography_paragraphs(doc: Document, *, required: bool = True) -> list:
    """Return only the reference block, stopping before a later appendix section.

    TMLR places references before appendices. Applying bibliography typography
    or counting numbered paragraphs to the rest of the document would corrupt
    both the appendix style and the data-derived reference count.
    """
    paragraphs = doc.paragraphs
    headings = [index for index, paragraph in enumerate(paragraphs)
                if paragraph.text.strip() == "References"]
    if not headings:
        if required:
            raise RuntimeError("References heading is missing from the DOCX")
        return []
    if len(headings) != 1:
        raise RuntimeError("expected one References heading in the DOCX")
    start = headings[0] + 1
    end = next(
        (index for index in range(start, len(paragraphs))
         if paragraphs[index].style is not None and paragraphs[index].style.name == "Heading 1"),
        len(paragraphs),
    )
    return paragraphs[start:end]


def postprocess(raw_docx: Path, output: Path) -> None:
    doc = Document(raw_docx)
    configure_styles(doc)
    normalize_compact_heading_hierarchy(doc)
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.74)
        section.right_margin = Inches(0.74)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        add_page_number(section)
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for shape, alt in zip(doc.inline_shapes, FIGURE_ALTS, strict=False):
        if shape.width > Inches(5.9):
            ratio = shape.height / shape.width
            shape.width = Inches(5.9)
            shape.height = int(shape.width * ratio)
        shape._inline.docPr.set("descr", alt)
        shape._inline.docPr.set("title", alt.split(".")[0])
    protect_numeric_citations_in_tables(doc)
    format_tables(doc)
    for paragraph in bibliography_paragraphs(doc, required=False):
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(8.2)
    doc.core_properties.title = "K-Bound: When Is Label-Free Adaptation Knowable?"
    doc.core_properties.subject = "Editable Word version of the final compact K-Bound manuscript"
    doc.core_properties.author = "Pratik Niroula"
    doc.core_properties.keywords = "test-time adaptation, distribution shift, K-Bound, KGA"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def count_docx_references(doc: Document) -> int:
    """Count numbered reference paragraphs, excluding later appendix lists."""
    return sum(
        1
        for paragraph in bibliography_paragraphs(doc)
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    )


def validate_docx(
    path: Path,
    reference_count: int,
    required_value_counts: Mapping[str, int] | None = None,
) -> None:
    doc = Document(path)
    if len(doc.inline_shapes) != 3:
        raise RuntimeError(f"expected 3 embedded figures, found {len(doc.inline_shapes)}")
    if len(doc.tables) < 13:
        raise RuntimeError(f"expected at least 13 manuscript tables, found {len(doc.tables)}")
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    for token in (r"\textsc", r"\begin{", r"\cite", "[eq:", "[fig:", "[tab:"):
        if token in xml:
            raise RuntimeError(f"unresolved conversion token remains in DOCX: {token}")
    for value, minimum_count in (required_value_counts or {}).items():
        observed_count = xml.lower().count(value.lower())
        if observed_count < minimum_count:
            raise RuntimeError(
                "generated provenance value was not preserved in the DOCX: "
                f"{value} (expected at least {minimum_count}, found {observed_count})"
            )
    if "iWildCam" not in xml or "withheld" not in xml:
        raise RuntimeError("iWildCam withholding language is missing from the DOCX")
    observed_references = count_docx_references(doc)
    if observed_references != reference_count:
        raise RuntimeError(
            "active bibliography count changed during conversion: "
            f"expected {reference_count}, found {observed_references}"
        )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output",
        type=Path,
        default=ROOT / "kbound_short_final_draft.docx",
        help="final DOCX path",
    )
    parser.add_argument("--keep-preprocessed", type=Path, help="optional path for the flattened TeX")
    args = parser.parse_args()
    latexpand = require_binary("latexpand")
    pandoc = require_binary("pandoc")
    macro_sources = discover_generated_macro_sources()
    macros = generated_macros(macro_sources)
    with tempfile.TemporaryDirectory(prefix="kbound-docx-") as temp_name:
        temp = Path(temp_name)
        flattened = temp / "kbound_submission_flat.tex"
        run_checked([latexpand, "--fatal", str(SOURCE)], cwd=ROOT, stdout_path=flattened)
        processed, reference_count, required_values = preprocess_with_metadata(
            flattened.read_text(encoding="utf-8"), macros=macros
        )
        preprocessed = temp / "kbound_submission_word.tex"
        preprocessed.write_text(processed, encoding="utf-8")
        if args.keep_preprocessed:
            args.keep_preprocessed.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(preprocessed, args.keep_preprocessed)
        raw_docx = temp / "kbound_short_raw.docx"
        resource_path = ":".join(
            str(path)
            for path in (
                ROOT,
                ROOT / "figures",
                ROOT / "paper",
                ROOT / "paper" / "generated",
                ROOT / "paper" / "sections",
            )
        )
        proc = subprocess.run(
            [
                pandoc,
                str(preprocessed),
                "-o",
                str(raw_docx),
                "--from",
                "latex",
                f"--resource-path={resource_path}",
                "--standalone",
            ],
            cwd=ROOT,
            text=True,
            capture_output=True,
            check=False,
        )
        if proc.returncode:
            raise RuntimeError(f"Pandoc failed ({proc.returncode}):\n{proc.stderr}")
        warnings = [line for line in proc.stderr.splitlines() if line.strip()]
        if warnings:
            raise RuntimeError("Pandoc emitted conversion warnings:\n" + "\n".join(warnings))
        postprocess(raw_docx, args.output.resolve())
    validate_docx(args.output.resolve(), reference_count, required_values)
    print(f"Built {args.output.resolve()}")


if __name__ == "__main__":
    main()
