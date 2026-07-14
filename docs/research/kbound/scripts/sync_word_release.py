#!/usr/bin/env python3
"""Synchronize the tracked Word draft with release-critical LaTeX corrections."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def replace_paragraph(document, prefix: str, text: str) -> None:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if not matches and any(paragraph.text == text for paragraph in document.paragraphs):
        return
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph beginning {prefix!r}; found {len(matches)}")
    matches[0].text = text


def remove_row(table, index: int) -> None:
    row = table.rows[index]
    table._tbl.remove(row._tr)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()

    document = Document(args.docx)
    replace_paragraph(
        document,
        "The evidence map uses only",
        (
            "The evidence map uses only unlabeled batches and model outputs. In the promoted "
            "per-condition artifacts retained by this release, it is an 11-dimensional vector "
            "computed from the frozen and adapted softmax outputs: mean predictive entropy and "
            "confidence; normalized marginal entropy; entropy and marginal-balance drops; the "
            "fraction of adapted predictions above confidence 0.9; the KL divergence between "
            "adapted and frozen predicted-class marginals; and the ℓ₂ norm of the parameter "
            "update. An exploratory 16-feature natural-shift panel is not used to support a "
            "promoted result. KGA estimates Δ directly from the retained evidence vector; it "
            "does not estimate the population margin M."
        ),
    )
    replace_paragraph(
        document,
        "The benefit model predicts",
        (
            "The benefit model predicts Δ_hat from labeled development conditions using a "
            "gradient-boosted regression tree (squared error, 250 trees, depth 2, learning rate "
            "0.05, subsample 0.8, seed 0), refit per dataset and candidate adapter. The complete "
            "promoted 11-feature schema is given in the appendix, with per-adapter update "
            "hyperparameters in the following table."
        ),
    )
    replace_paragraph(
        document,
        "Label-free evidence actually",
        (
            "Label-free evidence in the promoted per-condition artifacts. Tracks retained only "
            "as reconciled summaries do not support a stronger feature-level claim."
        ),
    )
    replace_paragraph(
        document,
        "KGA cannot avoid the cost",
        (
            "KGA cannot avoid the cost of proposing an update because candidate adaptation "
            "includes forward and backward computation before commitment. The controller-only "
            "artifact experiments/kbound/results/controller_cost_v1/cost_profile.json measures "
            "0.098 ms for evidence extraction and 0.245 ms for benefit-model evaluation "
            "(0.343 ms total) on Apple silicon with Python 3.12.13. The referenced "
            "ResNet-18/CIFAR checkpoint is 44.77 MB on disk and the serialized benefit model is "
            "194.9 KB. These environment-specific values exclude candidate adaptation and model "
            "inference."
        ),
    )
    replace_paragraph(
        document,
        "KGA controller added cost",
        (
            "KGA controller added cost (ResNet-18/CIFAR, Apple silicon, Python 3.12.13; "
            "scripts/recompute/cost_profile.py). Timings are environment-specific and exclude "
            "adaptation and inference."
        ),
    )
    next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("KGA controller added cost")
    ).paragraph_format.page_break_before = True
    replace_paragraph(
        document,
        "The one natural-shift track",
        (
            "The one natural-shift track with per-condition logs at multiple seeds is "
            "Camelyon17 (4 seeds, 9 composition-stress conditions per seed). The committed logs "
            "under experiments/kbound/results/camelyon17_multiseed_v1/ reproduce the "
            "candidate-dependent result: Tent is stable no-harm with FA_u=0 across seeds, EATA "
            "is inconclusive, and the helpful-dominated SAR arm over-freezes with FA_u up to "
            "0.11. The recomputation uses scripts/recompute/multiseed_natural.py. iWildCam, "
            "Office-Home, and RxRx1 multi-seed replays remain future work."
        ),
    )
    replace_paragraph(
        document,
        "Two label-free evidence panels",
        (
            "The promoted per-condition artifacts use the 11-feature label-free evidence panel "
            "implemented in kbound_pkg/kbound/evidence.py. The retained Camelyon17 multi-seed "
            "logs record the same ordered vector. Tracks available only through reconciled "
            "summary artifacts are not used to claim a richer feature schema."
        ),
    )
    replace_paragraph(document, "Base -dim evidence panel", "Promoted 11-dimensional evidence panel.")
    replace_paragraph(document, "@llp0.26@ Feature", "")

    evidence = next(
        table
        for table in document.tables
        if table.cell(0, 0).text == "Family" and "Frozen output" in table.cell(1, 0).text
    )
    evidence.cell(1, 1).text = "mean entropy, mean confidence, normalized marginal-class entropy"
    evidence.cell(2, 1).text = (
        "mean entropy, mean confidence, normalized marginal entropy, fraction confidence > 0.9"
    )
    evidence.cell(3, 1).text = "entropy drop, marginal-balance drop"
    if evidence.cell(4, 0).text == "Disagreement":
        remove_row(evidence, 4)
    evidence.cell(4, 0).text = "Distribution drift"
    evidence.cell(4, 1).text = "adapted-vs-frozen predicted-class marginal KL"
    evidence.cell(5, 0).text = "Update behavior"
    evidence.cell(5, 1).text = "ℓ₂ norm of the adapter parameter update"

    cost = next(
        table
        for table in document.tables
        if table.cell(0, 0).text == "Component" and "Evidence extraction" in table.cell(1, 0).text
    )
    cost.cell(1, 1).text = "0.098 ms"
    cost.cell(2, 1).text = "0.245 ms"
    cost.cell(3, 1).text = "0.343 ms"
    cost.cell(4, 0).text = "Rollback checkpoint (on-disk size)"
    cost.cell(4, 1).text = "44.77 MB"
    cost.cell(5, 0).text = "Serialized benefit model"
    cost.cell(5, 1).text = "194.9 KB"

    document.save(args.docx)


if __name__ == "__main__":
    main()
