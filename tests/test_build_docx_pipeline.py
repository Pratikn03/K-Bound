from __future__ import annotations

import hashlib
import importlib.util
from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "docs/research/kbound/scripts/build_docx.py"
SPEC = importlib.util.spec_from_file_location("kbound_build_docx", SCRIPT)
MODULE = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(MODULE)


def _write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def test_generated_sources_follow_live_dependency_graph(tmp_path: Path) -> None:
    source = tmp_path / "paper.tex"
    generated = tmp_path / "paper/generated"
    base = generated / "kbound_numbers.tex"
    cct = generated / "cct20_numbers.tex"
    _write(base, r"\newcommand{\BaseMetric}{0.10}")
    _write(cct, r"\newcommand{\CCTMetric}{0.20}")
    _write(source, r"\input{paper/generated/kbound_numbers}")

    sources = MODULE.discover_generated_macro_sources(source, generated_dir=generated, search_root=tmp_path)
    assert sources == (base.resolve(),)
    assert MODULE.generated_macros(sources) == {"BaseMetric": "0.10"}

    _write(
        source,
        "\n".join(
            (
                r"\input{paper/generated/kbound_numbers}",
                r"\input{paper/generated/cct20_numbers}",
            )
        ),
    )
    sources = MODULE.discover_generated_macro_sources(source, generated_dir=generated, search_root=tmp_path)
    assert sources == (base.resolve(), cct.resolve())
    assert MODULE.generated_macros(sources) == {
        "BaseMetric": "0.10",
        "CCTMetric": "0.20",
    }


def test_generated_macros_parse_nested_values_and_reject_conflicts(tmp_path: Path) -> None:
    first = tmp_path / "base_numbers.tex"
    second = tmp_path / "cct20_numbers.tex"
    _write(first, r"\newcommand{\Verdict}{\textnormal{not defined}}")
    _write(second, r"\newcommand{\Verdict}{complete}")

    assert MODULE.generated_macros((first,)) == {"Verdict": r"\textnormal{not defined}"}
    with pytest.raises(RuntimeError, match="defined by both"):
        MODULE.generated_macros((first, second))
    with pytest.raises(RuntimeError, match="no generated macro sources"):
        MODULE.generated_macros(())


def test_generated_macros_resolve_transitive_values_and_reject_cycles(
    tmp_path: Path,
) -> None:
    source = tmp_path / "numbers.tex"
    _write(
        source,
        "\n".join(
            (
                r"\newcommand{\ReleaseSHA}{\CanonicalSHA}",
                r"\newcommand{\CanonicalSHA}{\ArtifactSHA}",
                rf"\newcommand{{\ArtifactSHA}}{{{'a' * 64}}}",
            )
        ),
    )
    assert MODULE.generated_macros((source,)) == {
        "ReleaseSHA": "a" * 64,
        "CanonicalSHA": "a" * 64,
        "ArtifactSHA": "a" * 64,
    }

    _write(
        source,
        "\n".join(
            (
                r"\newcommand{\First}{\Second}",
                r"\newcommand{\Second}{\First}",
            )
        ),
    )
    with pytest.raises(RuntimeError, match="cyclic generated macro definition"):
        MODULE.generated_macros((source,))


def test_generated_sources_are_discovered_through_nested_inputs(tmp_path: Path) -> None:
    source = tmp_path / "paper.tex"
    body = tmp_path / "sections/body.tex"
    generated = tmp_path / "paper/generated"
    numbers = generated / "cct20_numbers.tex"
    _write(source, r"\input{sections/body}")
    _write(body, r"\input{paper/generated/cct20_numbers}")
    _write(numbers, r"\newcommand{\CCTMetric}{0.20}")

    assert MODULE.discover_generated_macro_sources(
        source,
        generated_dir=generated,
        search_root=tmp_path,
    ) == (numbers.resolve(),)


def test_hash_requirements_are_derived_only_from_used_macros() -> None:
    base_sha = "a" * 64
    cct_sha = "b" * 64
    tex = (
        r"\documentclass{article}"
        "\n"
        r"\begin{document}\BaseSHA, \BaseSHA, \CCTExecutionSHA\end{document}"
    )
    requirements = MODULE.generated_value_requirements(
        tex,
        {
            "BaseSHA": base_sha,
            "CCTExecutionSHA": cct_sha,
            "UnusedPendingSHA": r"\textnormal{pending}",
        },
    )
    assert requirements == {base_sha: 2, cct_sha: 1}

    with pytest.raises(RuntimeError, match="not a complete SHA-256"):
        MODULE.generated_value_requirements(
            tex,
            {"BaseSHA": base_sha, "CCTExecutionSHA": "pending"},
        )


def test_generated_definitions_in_document_body_are_removed_before_expansion() -> None:
    digest = "c" * 64
    tex = "\n".join(
        (
            r"\documentclass{article}",
            r"\begin{document}",
            r"\newcommand{\CCTCheckpointCount}{5}",
            rf"\newcommand{{\CCTInferenceSHA}}{{{digest}}}",
            r"\newcommand{\LocalFormatting}[1]{\textbf{#1}}",
            r"There are \CCTCheckpointCount{} checkpoints (\CCTInferenceSHA).",
            r"\end{document}",
        )
    )
    macros = {"CCTCheckpointCount": "5", "CCTInferenceSHA": digest}

    requirements = MODULE.generated_value_requirements(tex, macros)
    expanded = MODULE.expand_generated_macros(tex, macros)

    assert requirements == {digest: 1}
    assert r"\newcommand{5}{5}" not in expanded
    assert r"\newcommand{\CCTCheckpointCount}{5}" not in expanded
    assert rf"\newcommand{{\CCTInferenceSHA}}{{{digest}}}" not in expanded
    assert r"\newcommand{\LocalFormatting}[1]{\textbf{#1}}" in expanded
    assert f"There are 5{{}} checkpoints ({digest})." in expanded


def test_split_bibliography_accepts_data_derived_55th_cct_reference() -> None:
    entries = "\n".join(rf"\bibitem{{key{index}}} Reference {index}." for index in range(1, 56))
    tex = (
        r"Citations \cite{key1,key55}."
        "\n"
        r"\begin{thebibliography}{99}"
        "\n" + entries + "\n" + r"\end{thebibliography}"
    )

    rendered, bibliography = MODULE.split_bibliography(tex)
    assert len(bibliography) == 55
    assert bibliography[-1] == ("key55", "Reference 55.")
    resolved = MODULE.resolve_citations(rendered, bibliography)
    assert "Citations [1, 55]." in resolved
    assert resolved.count(r"\item ") == 55


def test_split_bibliography_accepts_shared_author_year_wrapper() -> None:
    tex = "\n".join(
        (
            r"See \cite{first,second}.",
            r"\begin{thebibliography}{99}",
            r"\KBbibitem{first}{First Author}{2025} First reference.",
            r"\KBbibitem{second}{Second et al.}{2026} Second reference.",
            r"\end{thebibliography}",
        )
    )

    rendered, bibliography = MODULE.split_bibliography(tex)

    assert bibliography == [
        ("first", "First reference."),
        ("second", "Second reference."),
    ]
    assert MODULE.resolve_citations(rendered, bibliography).startswith("See [1, 2].")


def test_citation_resolution_handles_notes_and_fails_closed() -> None:
    entries = [("first", "First reference."), ("second", "Second reference.")]
    tex = r"See \citep[compare][p.~2]{second,first}."
    assert MODULE.resolve_citations(tex, entries) == "See [1, 2]."

    with pytest.raises(RuntimeError, match="unknown citation key"):
        MODULE.resolve_citations(r"See \cite{missing}.", entries)
    with pytest.raises(RuntimeError, match="unsupported or unresolved"):
        MODULE.resolve_citations(r"See \citeauthor{first}.", entries)


def test_split_bibliography_rejects_multiple_active_lists() -> None:
    bibliography = (
        r"\begin{thebibliography}{1}\bibitem{key} Reference."
        r"\end{thebibliography}"
    )
    with pytest.raises(RuntimeError, match="expected one active compact bibliography"):
        MODULE.split_bibliography(bibliography + bibliography)


def test_live_source_manifest_macro_matches_live_manifest_file() -> None:
    sources = MODULE.discover_generated_macro_sources()
    macros = MODULE.generated_macros(sources)
    manifest = ROOT / "experiments/kbound/results/reconciled_panels_v1/source_manifest.json"
    assert macros["SourceManifestSHA"] == hashlib.sha256(manifest.read_bytes()).hexdigest()


def test_small_tables_keep_rows_together_for_word_pagination() -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=2)

    MODULE.format_tables(doc)

    assert all(
        paragraph.paragraph_format.keep_with_next is True
        for row in table.rows[:-1]
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    assert all(
        paragraph.paragraph_format.keep_with_next is not True
        for cell in table.rows[-1].cells
        for paragraph in cell.paragraphs
    )


def test_long_tables_keep_rows_intact_but_allow_breaks_between_body_rows() -> None:
    doc = Document()
    table = doc.add_table(rows=8, cols=2)
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.keep_with_next = True
        cant_split = MODULE.OxmlElement("w:cantSplit")
        cant_split.set(MODULE.qn("w:val"), "false")
        row._tr.get_or_add_trPr().append(cant_split)

    # Reapplying the formatter must not accumulate row-property elements.
    MODULE.format_tables(doc)
    MODULE.format_tables(doc)

    for row in table.rows:
        flags = row._tr.xpath("./w:trPr/w:cantSplit")
        assert len(flags) == 1
        assert flags[0].get(MODULE.qn("w:val")) == "true"
    assert len(table.rows[0]._tr.xpath("./w:trPr/w:tblHeader")) == 1
    assert all(
        paragraph.paragraph_format.keep_with_next is True
        for cell in table.rows[0].cells
        for paragraph in cell.paragraphs
    )
    assert all(
        paragraph.paragraph_format.keep_with_next is False
        for row in table.rows[1:]
        for cell in row.cells
        for paragraph in cell.paragraphs
    )


def test_table_caption_stays_with_table_without_binding_image_captions() -> None:
    doc = Document()
    doc.styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Image Caption", WD_STYLE_TYPE.PARAGRAPH)

    MODULE.configure_styles(doc)

    assert doc.styles["Table Caption"].paragraph_format.keep_with_next is True
    assert doc.styles["Image Caption"].paragraph_format.keep_with_next is not True
    assert doc.styles["Caption"].paragraph_format.keep_with_next is not True


def test_header_only_table_does_not_bind_following_prose() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=2)

    MODULE.format_tables(doc)

    assert len(table.rows[0]._tr.xpath("./w:trPr/w:cantSplit")) == 1
    assert all(
        paragraph.paragraph_format.keep_with_next is False
        for cell in table.rows[0].cells
        for paragraph in cell.paragraphs
    )


def test_reader_facing_track_tables_reserve_space_for_dataset_names() -> None:
    doc = Document()
    roles = doc.add_table(rows=1, cols=4)
    for cell, value in zip(
        roles.rows[0].cells,
        ("Track", "Shift geometry", "Candidate/protocol", "Defensible use"),
        strict=True,
    ):
        cell.text = value
    provenance = doc.add_table(rows=1, cols=4)
    for cell, value in zip(
        provenance.rows[0].cells,
        ("Track", "Released status", "Present evidence", "Remaining limitation"),
        strict=True,
    ):
        cell.text = value

    assert MODULE.table_widths(roles, 9960) == [2150, 2300, 2200, 3310]
    assert MODULE.table_widths(provenance, 9960) == [1800, 1600, 3100, 3460]


def test_numeric_table_citations_receive_word_joiners() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = "Track"
    table.rows[1].cells[0].text = "So2Sat-LCZ42\u00a0[30]"

    MODULE.protect_numeric_citations_in_tables(doc)

    assert table.rows[1].cells[0].text == "So2Sat-LCZ42\u00a0[\u206030\u2060]"


def test_section_numbering_counts_unlabelled_headings_and_renders_numbers() -> None:
    tex = "\n".join(
        (
            r"\section{Introduction}\label{sec:intro}",
            r"\subsection{First subsection}",
            r"\subsection{Second subsection}\label{sec:second}",
            r"\subsubsection{First detail}",
            r"\subsubsection{Second detail}\label{sec:detail}",
            r"\section{Unlabelled main section}",
            r"\section{Results}\label{sec:results}",
        )
    )

    numbered, labels = MODULE.number_section_headings(tex)

    assert labels == {
        "sec:intro": "1",
        "sec:second": "1.2",
        "sec:detail": "1.2.2",
        "sec:results": "3",
    }
    assert r"\subsection{1.1 First subsection}" in numbered
    assert r"\subsection{1.2 Second subsection}" in numbered
    assert r"\subsubsection{1.2.2 Second detail}" in numbered
    assert r"\section{2 Unlabelled main section}" in numbered
    assert r"\section{3 Results}" in numbered
    assert MODULE.section_labels(tex) == labels


def test_section_numbering_resets_appendices_and_preserves_starred_headings() -> None:
    tex = "\n".join(
        (
            r"\section{Main}\label{sec:main}",
            r"\subsection{Main detail}",
            r"\appendix",
            r"\section*{Appendix guide}",
            r"\section{First appendix}\label{app:first}",
            r"\subsection*{Unnumbered note}",
            r"\subsection{Detail}\label{app:detail}",
            r"\subsubsection{Subdetail}\label{app:subdetail}",
            r"\section{Second appendix}\label{app:second}",
            r"\subsection{Reset detail}\label{app:reset}",
            r"\section*{References}",
        )
    )

    numbered, labels = MODULE.number_section_headings(tex)

    assert labels == {
        "sec:main": "1",
        "app:first": "A",
        "app:detail": "A.1",
        "app:subdetail": "A.1.1",
        "app:second": "B",
        "app:reset": "B.1",
    }
    assert r"\section{A First appendix}" in numbered
    assert r"\subsection{B.1 Reset detail}" in numbered
    for original in (
        r"\section*{Appendix guide}",
        r"\subsection*{Unnumbered note}",
        r"\section*{References}",
    ):
        assert original in numbered


def test_section_numbering_handles_comments_nested_titles_and_short_titles() -> None:
    tex = "\n".join(
        (
            r"% \section{Commented-out heading}\label{sec:ignored}",
            r"\section[Short {A]}]{Rate \% and \textbf{evidence}\label{sec:title}}",
            r"% A comment between the heading and its second label.",
            r"\label{sec:alias}",
            r"\subsection % The title starts on the next line.",
            r"{Bridge} % \subsection{Another ignored heading}",
            r"\label{sec:bridge}",
        )
    )

    numbered, labels = MODULE.number_section_headings(tex)

    assert labels == {"sec:title": "1", "sec:alias": "1", "sec:bridge": "1.1"}
    assert r"\section[Short {A]}]{1 Rate \% and \textbf{evidence}" in numbered
    assert "{1.1 Bridge}" in numbered
    assert r"% \section{Commented-out heading}" in numbered


@pytest.mark.parametrize(
    ("tex", "message"),
    (
        (r"\section Unbraced", "literal braced title"),
        (r"\section{Unclosed", "unbalanced title"),
        (r"\section[Unclosed", "unbalanced optional title"),
        (
            r"\section{First}\label{duplicate}\section{Second}\label{duplicate}",
            "duplicate section label",
        ),
        (r"\appendix\subsection{No appendix section}", "article-style section"),
    ),
)
def test_section_numbering_rejects_ambiguous_headings(tex: str, message: str) -> None:
    with pytest.raises(RuntimeError, match=message):
        MODULE.number_section_headings(tex)


def test_section_numbering_matches_article_appendix_limit() -> None:
    tex = r"\appendix" + "".join(
        rf"\section{{Appendix {index}}}\label{{app:{index}}}" for index in range(1, 27)
    )
    numbered, labels = MODULE.number_section_headings(tex)
    assert labels["app:26"] == "Z"
    assert r"\section{Z Appendix 26}" in numbered
    with pytest.raises(RuntimeError, match="article-style section"):
        MODULE.number_section_headings(tex + r"\section{Too many appendices}")


def test_heading_crossrefs_preserve_equation_theorem_and_float_numbers() -> None:
    tex = "\n".join(
        (
            r"\section{Main}\label{sec:main}",
            r"\subsection{Unlabelled}",
            r"\subsection{Result}\label{sec:result}",
            r"\begin{theorem}\label{thm:first}",
            r"\begin{equation}x=1.\label{eq:first}\end{equation}",
            r"\end{theorem}",
            r"\begin{table*}\caption{Values.}\label{tab:first}\end{table*}",
            r"\begin{figure}\caption{Diagram.}\label{fig:first}\end{figure}",
            r"\appendix\section{Evidence}\label{app:evidence}",
            r"Section~\ref{sec:result}, Appendix~\ref{app:evidence},",
            r"Theorem~\ref{thm:first}, Eq.~\eqref{eq:first},",
            r"Table~\ref{tab:first}, and Figure~\ref{fig:first}.",
        )
    )

    resolved = MODULE.resolve_cross_references(tex)

    assert r"\subsection{1.2 Result}" in resolved
    assert r"\section{A Evidence}" in resolved
    assert "Section~1.2, Appendix~A," in resolved
    assert "Theorem~1, Eq.~(1)," in resolved
    assert "Table~1, and Figure~1." in resolved
    assert r"\caption{Table 1. Values.}" in resolved
    assert r"\caption{Figure 1. Diagram.}" in resolved
    assert resolved.count(r"\text{(1)}") == 1
    assert r"\label{" not in resolved
    assert r"\ref{" not in resolved


def test_references_to_starred_headings_fail_closed() -> None:
    tex = r"\section*{Unnumbered}\label{sec:starred}See~\ref{sec:starred}."
    with pytest.raises(RuntimeError, match="unresolved reference: sec:starred"):
        MODULE.resolve_cross_references(tex)


def test_release_section_numbers_match_the_maintained_article_hierarchy() -> None:
    body = ROOT / "docs/research/kbound/kbound_submission_body.tex"
    supplement = ROOT / "docs/research/kbound/kbound_submission_supplement.tex"
    labels = MODULE.section_labels(body.read_text(encoding="utf-8") + "\n" + supplement.read_text(encoding="utf-8"))

    assert labels["sec:compact-bridge"] == "6.1"
    assert labels["sec:so2sat-development-stop"] == "8.4"
    assert labels["app:population-transfer"] == "A"
    assert labels["app:compact-confirmation"] == "K"


def test_word_theory_notation_preserves_overbar_and_roman_clauses() -> None:
    tex = (
        r"$\bar a+\bar{a}+\bar a_\theta+\bar b$"
        r"\begin{enumerate}"
        r"\item[\textnormal{(i)}] Reduction."
        r"\item[\textnormal{(ii)}] Sufficiency."
        r"\item[\textnormal{(iii)}] Necessity."
        r"\item[\textnormal{(iv)}] Strict commitment."
        r"\end{enumerate}"
        r"Proof. Part (i) and part (iv)."
        r"\begin{enumerate}\item Ordinary item.\end{enumerate}"
    )
    converted = MODULE.normalize_word_theory_notation(tex)
    assert converted.count(r"\overline{a}") == 3
    assert r"\bar b" in converted
    assert r"\begin{enumerate}[(i)]" in converted
    assert r"\item[\textnormal" not in converted
    assert "Proof. Part (i) and part (iv)." in converted
    assert r"\begin{enumerate}\item Ordinary item.\end{enumerate}" in converted
    assert MODULE.normalize_word_theory_notation(converted) == converted


def test_word_theory_notation_rejects_incomplete_roman_list() -> None:
    tex = r"\begin{enumerate}\item[\textnormal{(i)}] Only one.\end{enumerate}"
    with pytest.raises(RuntimeError, match="four consecutive Roman theorem clauses"):
        MODULE.normalize_word_theory_notation(tex)


def test_word_figure_alt_text_uses_residual_not_drift() -> None:
    assert "calibration-residual bound beta" in MODULE.FIGURE_ALTS[1]
    assert "drift budget" not in MODULE.FIGURE_ALTS[1]


def _numbered_paragraph(doc, text):
    paragraph = doc.add_paragraph(text)
    numbering = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    numbering.get_or_add_numId().val = 7
    numbering.get_or_add_ilvl().val = 0
    return paragraph


def test_reference_count_excludes_numbered_appendix_lists():
    doc = Document()
    doc.add_heading("References", level=1)
    _numbered_paragraph(doc, "First reference.")
    _numbered_paragraph(doc, "Second reference.")
    doc.add_heading("A Supplementary material", level=1)
    _numbered_paragraph(doc, "Not a third reference.")

    assert MODULE.count_docx_references(doc) == 2
    assert [p.text for p in MODULE.bibliography_paragraphs(doc)] == [
        "First reference.", "Second reference."
    ]


def test_bibliography_typography_stops_before_appendix(tmp_path):
    raw = tmp_path / "raw.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_heading("References", level=1)
    _numbered_paragraph(doc, "A published reference.")
    doc.add_heading("A Population transfer", level=1)
    doc.add_paragraph("Appendix body retains ordinary body typography.")
    doc.save(raw)

    MODULE.postprocess(raw, output)

    rendered = Document(output)
    reference = next(p for p in rendered.paragraphs if p.text == "A published reference.")
    appendix = next(p for p in rendered.paragraphs if p.text.startswith("Appendix body"))
    # OOXML serializes font sizes in half-points.
    assert reference.runs[0].font.size == MODULE.Pt(8.0)
    assert appendix.runs[0].font.size is None
    assert appendix.style.font.size == MODULE.Pt(9.5)


def test_reference_counter_still_accepts_terminal_bibliography():
    doc = Document()
    doc.add_heading("References", level=1)
    _numbered_paragraph(doc, "Only reference.")
    assert MODULE.count_docx_references(doc) == 1


def test_reference_counter_rejects_missing_or_duplicate_heading():
    doc = Document()
    with pytest.raises(RuntimeError, match="References heading is missing"):
        MODULE.count_docx_references(doc)
    assert MODULE.bibliography_paragraphs(doc, required=False) == []
    doc.add_heading("References", level=1)
    doc.add_heading("References", level=1)
    with pytest.raises(RuntimeError, match="expected one References heading"):
        MODULE.count_docx_references(doc)


def test_numbered_align_rows_preserve_equation_references_and_counter():
    source = (
        r"\begin{equation}a=0\label{eq:before}\end{equation}"
        r"\begin{align}b&=1\label{eq:first}\\c&=2\label{eq:second}\end{align}"
        r"See \eqref{eq:before}, \eqref{eq:first}, \eqref{eq:second}."
    )
    flattened = MODULE.simplify_alignment_math(source)
    assert flattened.count(r"\begin{equation}") == 3
    resolved = MODULE.resolve_cross_references(flattened)
    assert "See (1), (2), (3)." in resolved
    assert r"\label{" not in resolved
    assert r"\eqref{" not in resolved


def test_starred_align_does_not_create_numbered_equations():
    source = r"\begin{align*}a&=0\\b&=1\end{align*}"
    flattened = MODULE.simplify_alignment_math(source)
    assert r"\begin{equation}" not in flattened
    assert flattened.count(r"\[") == 2


def test_align_nonumber_row_does_not_advance_the_counter():
    source = (
        r"\begin{align}a&=0\nonumber\\b&=1\label{eq:numbered}\end{align}"
        r"See \eqref{eq:numbered}."
    )
    flattened = MODULE.simplify_alignment_math(source)
    assert flattened.count(r"\begin{equation}") == 1
    assert "See (1)." in MODULE.resolve_cross_references(flattened)


def _word_algorithm_source() -> str:
    return r"""\begin{algorithm}[t]
\caption{KGA for one proposed update}
\label{alg:compact-kga}
\begin{algorithmic}[1]
\Require Frozen $f_0$, adapter $\mathcal A$, unlabeled $X$, frozen $h_\theta$ and calibration artifacts
\State Create $f_a\leftarrow\mathcal A(f_0,X)$ in isolated shadow state
\State Extract $Z=\phi(X,f_0,f_a)$ and validate the assessment inputs
\If{the assessment is unavailable}
  \State retain $f_0$, log the reason, and return \abstain
\EndIf
\State Obtain $\varepsilon$ from the calibrated residual rule in Eq.~\eqref{eq:compact-rank}
\State $\widehat\Delta\leftarrow h_\theta(Z)$; $L\leftarrow\widehat\Delta-\varepsilon$; $U\leftarrow\widehat\Delta+\varepsilon$
\If{$L>0$} \State commit $f_a$ and return \adapt
\ElsIf{$U<0$} \State discard $f_a$ and return \freeze
\Else \State retain $f_0$, log uncertainty, and return \abstain
\EndIf
\end{algorithmic}
\end{algorithm}"""


def test_word_algorithm_preserves_unavailable_assessment_before_radius():
    converted = MODULE.replace_algorithm_for_word(_word_algorithm_source())

    validation = converted.index("validate the assessment inputs")
    unavailable = converted.index("the assessment is unavailable")
    fallback = converted.index(r"retain $f_0$, log the reason, and return \abstain")
    radius = converted.index(r"Obtain $\varepsilon$ from the calibrated residual rule")
    prediction = converted.index(r"$\widehat\Delta\leftarrow h_\theta(Z)$")
    assert validation < unavailable < fallback < radius < prediction
    assert r"\eqref{eq:compact-rank}" in converted
    assert converted.count(r"\abstain") == 2


def test_word_algorithm_derives_editable_steps_and_caption_from_source():
    source = _word_algorithm_source().replace(
        "KGA for one proposed update", r"Updated \textbf{source} caption"
    ).replace(
        r"\State Obtain", "\\State Record the revised review reason\n\\State Obtain"
    ).replace("calibration artifacts", "versioned residual artifacts")
    converted = MODULE.replace_algorithm_for_word("Before.\n" + source + "\nAfter.")

    assert r"\paragraph{Algorithm 1. Updated \textbf{source} caption}" in converted
    assert r"\item Record the revised review reason" in converted
    assert "versioned residual artifacts" in converted
    assert converted.count(r"\begin{enumerate}") == converted.count(r"\end{enumerate}")
    assert r"\begin{algorithm}" not in converted
    assert r"\begin{algorithmic}" not in converted
    assert converted.startswith("Before.\n")
    assert converted.endswith("\nAfter.")


def test_word_algorithm_keeps_strict_if_elseif_else_actions():
    converted = MODULE.replace_algorithm_for_word(_word_algorithm_source())

    assert r"\item \textbf{If} $L>0$:" in converted
    assert r"\item commit $f_a$ and return \adapt" in converted
    assert r"\item \textbf{Else if} $U<0$:" in converted
    assert r"\item discard $f_a$ and return \freeze" in converted
    assert r"\item \textbf{Else}:" in converted
    assert r"\item retain $f_0$, log uncertainty, and return \abstain" in converted


@pytest.mark.parametrize(
    "unsupported",
    (
        r"\For{each candidate}",
        r"\While{waiting}",
        r"\Statex Additional instruction",
        r"\UnknownStep{extra action}",
        r"\State apply \UnknownAction",
    ),
)
def test_word_algorithm_rejects_unsupported_commands_instead_of_dropping_them(unsupported):
    source = _word_algorithm_source().replace(
        r"\end{algorithmic}", unsupported + "\n" + r"\end{algorithmic}"
    )
    with pytest.raises(RuntimeError, match="unsupported algorithm"):
        MODULE.replace_algorithm_for_word(source)


@pytest.mark.parametrize(
    ("before", "after"),
    (
        (r"\If{the assessment is unavailable}", r"\Else"),
        (r"\EndIf", ""),
        (r"\Else \State", r"\Else \Else \State"),
        (r"\If{$L>0$}", r"\If{$L>0$"),
        (r"\If{the assessment is unavailable}", r"\If{}"),
    ),
)
def test_word_algorithm_rejects_malformed_control_flow(before, after):
    source = _word_algorithm_source().replace(before, after, 1)
    with pytest.raises(RuntimeError, match="algorithm"):
        MODULE.replace_algorithm_for_word(source)


@pytest.mark.parametrize("source", ("No algorithm.", _word_algorithm_source() * 2))
def test_word_algorithm_requires_one_source_algorithm(source):
    with pytest.raises(RuntimeError, match="expected one KGA algorithm"):
        MODULE.replace_algorithm_for_word(source)


def test_word_algorithm_current_manuscript_retains_both_abstention_paths():
    source = (ROOT / "docs/research/kbound/kbound_submission_body.tex").read_text(encoding="utf-8")
    converted = MODULE.replace_algorithm_for_word(source)

    assert r"\item retain $f_0$, log the reason, and return \abstain" in converted
    assert r"\item retain $f_0$, log uncertainty, and return \abstain" in converted
    assert r"\item Obtain $\varepsilon$ from the calibrated residual rule" in converted
